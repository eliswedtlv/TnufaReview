import io
import json

import pytest
from docx import Document

import app as app_module


DESIRED_ORDER = [
    "executive_summary",
    "the_need",
    "the_product",
    "team_and_capabilities",
    "intellectual_property",
    "technology_uniqueness_innovation",
    "tasks_and_activities",
    "market_clients_competition_business_model",
    "grant_contribution_to_success",
    "royalties",
    "economic_and_technological_contribution",
]


class _FakeMessage:
    def __init__(self, content):
        self.content = content


class _FakeChoice:
    def __init__(self, content):
        self.message = _FakeMessage(content)


class _FakeResponse:
    def __init__(self, content):
        self.choices = [_FakeChoice(content)]


def _make_docx_bytes():
    """A plain (non-Tnufa) docx — no Heading 1 anchors, no header marker."""
    doc = Document()
    doc.add_paragraph("סיכום מנהלים")
    doc.add_paragraph("הזן טקסט כאן...")  # placeholder — must be skipped
    doc.add_paragraph("זוהי תשובת היזם בפסקה")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "תא ראשון"
    table.rows[0].cells[1].text = "תא שני"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_realistic_form():
    """
    A realistic filled Tnufa form as real exports look: NO heading styles at
    all — only plain paragraphs plus a 2×1 table whose second cell holds the
    executive-summary answer. There are no Heading 1 anchors to key off, so the
    AI extractor is the only viable path.
    """
    doc = Document()
    doc.add_paragraph("בקשת השקעה מקרן תנופה")  # plain paragraph, not a heading
    doc.add_paragraph("סיכום מנהלים")
    doc.add_paragraph("הזן טקסט כאן...")  # placeholder — must be skipped

    # Exec-summary answer lives inside the 2nd cell of a 2×1 table.
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "סיכום מנהלים"
    table.rows[0].cells[1].text = "תשובת הסיכום של היזם בתוך תא טבלה"

    doc.add_paragraph("הצורך")
    doc.add_paragraph("תיאור הצורך שהמיזם פותר")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- extract_all_text_from_docx --------------------------------------------

def test_extract_all_text_reads_paragraphs_and_cells_in_order_skipping_placeholder():
    text = app_module.extract_all_text_from_docx(_make_docx_bytes())
    lines = text.split("\n")

    assert "הזן טקסט כאן..." not in lines  # placeholder skipped
    assert lines == [
        "סיכום מנהלים",
        "זוהי תשובת היזם בפסקה",
        "תא ראשון",
        "תא שני",
    ]


# --- ai_extract_sections (fallback) ----------------------------------------

def test_ai_extract_sections_coerces_to_exactly_11_string_keys(monkeypatch):
    bad_output = {
        "executive_summary": "סיכום",
        "the_need": 123,              # non-string -> coerced to str
        "the_product": None,          # non-string -> coerced to str
        "totally_unexpected": "junk",  # extra key -> ignored
        # remaining canonical keys missing -> filled with ""
    }

    def fake_create(**kwargs):
        return _FakeResponse(json.dumps(bad_output))

    monkeypatch.setattr(app_module.client.chat.completions, "create", fake_create)

    sections = app_module.ai_extract_sections("some text")

    assert set(sections.keys()) == set(DESIRED_ORDER)
    for key in DESIRED_ORDER:
        assert isinstance(sections[key]["applicant_answer"], str)
        assert sections[key]["question"] == app_module.SECTION_STRUCTURE[key]["question"]

    assert sections["executive_summary"]["applicant_answer"] == "סיכום"
    assert sections["the_need"]["applicant_answer"] == "123"
    assert sections["the_product"]["applicant_answer"] == "None"
    assert sections["intellectual_property"]["applicant_answer"] == ""  # missing -> ""


# --- realistic form without heading styles ---------------------------------

def test_extract_all_text_captures_exec_summary_table_cell():
    """Real forms have no heading styles; the exec-summary answer lives in a
    table cell. extract_all_text_from_docx must still capture it."""
    text = app_module.extract_all_text_from_docx(_make_realistic_form())
    lines = text.split("\n")

    assert "הזן טקסט כאן..." not in lines  # placeholder skipped
    assert "תשובת הסיכום של היזם בתוך תא טבלה" in lines  # cell text captured
    assert "תיאור הצורך שהמיזם פותר" in lines


# --- order_review coercion --------------------------------------------------

def test_order_review_coerces_missing_extra_and_malformed_keys():
    raw = {
        "executive_summary": ["הערה 1", "הערה 2"],
        "the_need": [],                 # empty -> sentinel
        "the_product": "מחרוזת בודדת",   # string -> wrapped
        "team_and_capabilities": [1, 2],  # non-string items -> str()
        "totally_unexpected": ["junk"],  # extra key -> dropped
        # remaining keys missing -> sentinel
    }
    ordered = app_module.order_review(raw)

    assert list(ordered.keys()) == DESIRED_ORDER
    assert ordered["executive_summary"] == ["הערה 1", "הערה 2"]
    assert ordered["the_need"] == ["אין הערות"]
    assert ordered["the_product"] == ["מחרוזת בודדת"]
    assert ordered["team_and_capabilities"] == ["1", "2"]
    assert ordered["intellectual_property"] == ["אין הערות"]
    assert "totally_unexpected" not in ordered


# --- /review endpoint -------------------------------------------------------

@pytest.fixture
def client():
    app_module.app.config["TESTING"] = True
    return app_module.app.test_client()


def test_review_ai_extract_then_single_review_call(client, monkeypatch):
    calls = {"extract": 0, "review": 0, "payload": None}

    def fake_extract(full_text):
        calls["extract"] += 1
        return {
            k: {
                "question": app_module.SECTION_STRUCTURE[k]["question"],
                "applicant_answer": "x",
            }
            for k in DESIRED_ORDER
        }

    def fake_create(**kwargs):
        calls["review"] += 1
        calls["payload"] = kwargs["messages"][-1]["content"]
        body = {k: ["אין הערות"] for k in DESIRED_ORDER}
        return _FakeResponse(json.dumps(body, ensure_ascii=False))

    monkeypatch.setattr(app_module, "ai_extract_sections", fake_extract)
    monkeypatch.setattr(app_module.client.chat.completions, "create", fake_create)

    resp = client.post(
        "/review",
        data={"file": (io.BytesIO(_make_realistic_form()), "form.docx")},
        content_type="multipart/form-data",
    )

    assert resp.status_code == 200
    # AI extraction is the sole path → one extract, exactly one review call.
    assert calls["extract"] == 1
    assert calls["review"] == 1

    # Payload contract: instructions_form + all 11 in sections_to_review, and
    # application_form uses applicant_answer.
    payload = calls["payload"]
    assert "instructions_form" in payload
    assert "general_instructions" in payload
    assert "sections_to_review" in payload
    assert "application_form" in payload
    assert "applicant_answer" in payload
    for key in DESIRED_ORDER:
        assert key in payload

    data = resp.get_json()
    assert set(data.keys()) == set(DESIRED_ORDER)
    assert len(data) == 11


def test_review_returns_all_11_keys(client, monkeypatch):
    def fake_create(**kwargs):
        user_content = kwargs["messages"][-1]["content"]
        if "application_form" in user_content:  # review call
            body = {k: ["אין הערות"] for k in DESIRED_ORDER}
        else:  # extraction call
            body = {k: f"answer-{k}" for k in DESIRED_ORDER}
        return _FakeResponse(json.dumps(body, ensure_ascii=False))

    monkeypatch.setattr(app_module.client.chat.completions, "create", fake_create)

    resp = client.post(
        "/review",
        data={"file": (io.BytesIO(_make_docx_bytes()), "form.docx")},
        content_type="multipart/form-data",
    )

    assert resp.status_code == 200
    data = resp.get_json()
    assert set(data.keys()) == set(DESIRED_ORDER)
    assert len(data) == 11


def test_review_rejects_non_docx_with_400(client):
    resp = client.post(
        "/review",
        data={"file": (io.BytesIO(b"this is not a docx"), "bad.docx")},
        content_type="multipart/form-data",
    )

    assert resp.status_code == 400
    assert "readable Word .docx" in resp.get_json()["error"]


def test_health_returns_ok(client):
    resp = client.get("/health")
    assert resp.status_code == 200
    assert resp.get_json() == {"status": "ok"}


def test_favicon_ico_returns_200(client):
    resp = client.get("/favicon.ico")
    assert resp.status_code == 200
    assert resp.mimetype == "image/x-icon"


# --- upload size cap --------------------------------------------------------

def test_max_content_length_is_10mb():
    assert app_module.app.config["MAX_CONTENT_LENGTH"] == 10 * 1024 * 1024


def test_review_rejects_oversized_upload_with_json_413(client):
    resp = client.post(
        "/review",
        data={"file": (io.BytesIO(b"x" * (11 * 1024 * 1024)), "big.docx")},
        content_type="multipart/form-data",
    )

    assert resp.status_code == 413
    assert "error" in resp.get_json()
